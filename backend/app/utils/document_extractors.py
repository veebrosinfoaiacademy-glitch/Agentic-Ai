"""Text extraction from uploaded documents.

One function per format, each taking raw bytes and returning
(text, metadata). Every extractor works entirely in memory — none of them
writes a temporary file, so there is no cleanup path that can leak or fail.

SECURITY: these functions parse bytes. They do not execute, import, compile
or evaluate anything, they do not run DOCX macros, and they do not follow
URLs found inside a document. An uploaded file is hostile input that gets
read and turned into a string, and nothing else.
"""

import csv
import io
import logging

from app.utils.errors import AppError

logger = logging.getLogger("app.documents")

# Read as bytes and decoded by us rather than by the parser, so a decoding
# problem surfaces as a clean error instead of mojibake.
_TEXT_ENCODING = "utf-8-sig"  # tolerates a UTF-8 BOM and strips it

# File signatures. Checked before parsing so a .pdf that is secretly a ZIP is
# rejected on evidence rather than on the client's say-so.
PDF_SIGNATURE = b"%PDF-"
ZIP_SIGNATURE = b"PK\x03\x04"

# CSV cells longer than this are truncated in the text rendering. Guards
# against a single enormous cell dominating the extracted output.
MAX_CSV_CELL_CHARS = 500


def _invalid(reason: str) -> AppError:
    """The file is not what it claims to be, or is corrupt."""
    logger.warning("Document rejected: %s", reason)
    return AppError(
        code="DOCUMENT_INVALID",
        message=f"The file could not be read: {reason}",
        status_code=422,
    )


def _extraction_failed(fmt: str, exc: Exception) -> AppError:
    """The parser raised something we did not anticipate.

    The exception type is logged; the client gets a generic message, because
    a parser traceback describes our internals, not their file.
    """
    logger.error("%s extraction failed (%s): %s", fmt, type(exc).__name__, exc)
    return AppError(
        code="DOCUMENT_EXTRACTION_FAILED",
        message="The document could not be processed.",
        status_code=422,
    )


# --- Normalisation ----------------------------------------------------------


def normalize_text(text: str) -> str:
    """Clean up extracted text without destroying its structure.

    Line endings are unified, null bytes dropped, and runs of blank lines
    collapsed to at most one. Paragraph and heading separation survives —
    stripping all blank lines would merge a document into a single wall of
    text and lose the structure a summarizer relies on.
    """
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\x00", "")

    # Trim trailing spaces per line, then collapse 3+ newlines into 2.
    lines = [line.rstrip() for line in text.split("\n")]

    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if line:
            blank_run = 0
            cleaned.append(line)
        else:
            blank_run += 1
            if blank_run <= 1:
                cleaned.append("")

    return "\n".join(cleaned).strip()


def _decode(data: bytes, fmt: str) -> str:
    """Decode bytes as UTF-8, failing loudly rather than corrupting text."""
    try:
        return data.decode(_TEXT_ENCODING)
    except UnicodeDecodeError as exc:
        logger.warning("%s is not valid UTF-8 at byte %d", fmt, exc.start)
        raise AppError(
            code="DOCUMENT_INVALID",
            message="The file is not valid UTF-8 text.",
            status_code=422,
        ) from None


# --- Extractors -------------------------------------------------------------


def extract_txt(data: bytes) -> tuple[str, dict]:
    """Plain text. Decoded, BOM-stripped and normalised."""
    return normalize_text(_decode(data, "TXT")), {"encoding": "utf-8"}


def extract_markdown(data: bytes) -> tuple[str, dict]:
    """Markdown, treated as text.

    The syntax is preserved deliberately — `## Heading` and `- bullet` are
    signal about document structure, and converting to HTML or stripping the
    markers would throw that away before the Content Agent ever sees it.
    """
    text = normalize_text(_decode(data, "Markdown"))
    heading_count = sum(1 for line in text.split("\n") if line.lstrip().startswith("#"))
    return text, {"encoding": "utf-8", "heading_count": heading_count}


def extract_csv(data: bytes) -> tuple[str, dict]:
    """CSV rendered as readable text rather than dumped as Python objects.

    Quoting and embedded commas are handled by the csv module. Cell values
    are only ever treated as text — a cell beginning with "=" is a formula
    to a spreadsheet, but here it is just a string, and nothing evaluates it.
    """
    text_data = _decode(data, "CSV")
    if not text_data.strip():
        return "", {"rows": 0, "columns": 0}

    try:
        reader = csv.reader(io.StringIO(text_data))
        rows = [row for row in reader if any(cell.strip() for cell in row)]
    except csv.Error as exc:
        raise _invalid(f"malformed CSV ({exc})") from None

    if not rows:
        return "", {"rows": 0, "columns": 0}

    def render(cell: str) -> str:
        cell = cell.strip()
        return cell[:MAX_CSV_CELL_CHARS] + "..." if len(cell) > MAX_CSV_CELL_CHARS else cell

    header, *data_rows = rows
    parts = [f"Header: {', '.join(render(c) for c in header)}"]

    if data_rows:
        parts.append("")
        parts.append("Rows:")
        # Ragged rows are kept as-is rather than dropped; a short row is
        # usually missing data, which is itself worth showing.
        parts.extend(" | ".join(render(c) for c in row) for row in data_rows)

    metadata = {
        "rows": len(data_rows),
        "columns": max(len(row) for row in rows),
        "has_header": True,
    }
    return normalize_text("\n".join(parts)), metadata


def extract_pdf(data: bytes) -> tuple[str, dict]:
    """PDF text, page by page.

    Uses pypdf, already part of the project's fixed stack. No OCR is
    performed: a scanned or image-only PDF yields no text, and that is
    reported honestly rather than papered over.
    """
    if not data.startswith(PDF_SIGNATURE):
        raise _invalid("not a PDF file")

    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))

        # An encrypted PDF may open but yield nothing. Try the empty password
        # (common for "owner password" PDFs) and give up cleanly otherwise.
        if reader.is_encrypted:
            try:
                if reader.decrypt("") == 0:
                    raise _invalid("the PDF is password protected")
            except AppError:
                raise
            except Exception:
                raise _invalid("the PDF is password protected") from None

        pages: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                pages.append(f"--- Page {index} ---\n{page_text}")

        metadata = {"page_count": len(reader.pages), "pages_with_text": len(pages)}

    except AppError:
        raise
    except PdfReadError as exc:
        raise _invalid(f"corrupt PDF ({exc})") from None
    except Exception as exc:
        raise _extraction_failed("PDF", exc) from None

    return normalize_text("\n\n".join(pages)), metadata


def extract_docx(data: bytes) -> tuple[str, dict]:
    """DOCX paragraphs and tables, in document order.

    python-docx exposes `.paragraphs` and `.tables` as separate lists, which
    loses their relative position — a table between two paragraphs would come
    out at the end. Walking the body XML instead keeps the reading order.

    Only text is read. Embedded objects, macros and external references are
    ignored, never executed and never fetched.
    """
    if not data.startswith(ZIP_SIGNATURE):
        raise _invalid("not a DOCX file")

    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        document = Document(io.BytesIO(data))
        body = document.element.body

        blocks: list[str] = []
        paragraph_count = 0
        table_count = 0

        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                paragraph_count += 1
                text = Paragraph(child, document).text.strip()
                if text:
                    blocks.append(text)

            elif child.tag == qn("w:tbl"):
                table_count += 1
                rows = []
                for row in Table(child, document).rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    blocks.append("\n".join(rows))

        metadata = {"paragraph_count": paragraph_count, "table_count": table_count}

    except AppError:
        raise
    except Exception as exc:
        # python-docx raises PackageNotFoundError for a non-Word zip, plus a
        # variety of XML errors for corrupt ones.
        if "PackageNotFound" in type(exc).__name__:
            raise _invalid("not a valid Word document") from None
        raise _extraction_failed("DOCX", exc) from None

    return normalize_text("\n\n".join(blocks)), metadata
