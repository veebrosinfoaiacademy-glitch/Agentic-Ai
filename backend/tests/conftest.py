"""Shared test fixtures.

The fake client here lets every database test run offline. No test in this
suite opens a network connection, so the suite works for anyone who clones
the repository without a MongoDB Atlas account.
"""

import io
from collections.abc import Iterator

import pytest
from pymongo.errors import ConnectionFailure

from app.config import settings
from app.database import mongodb
from app.services.groq_service import GroqResult, groq_service
from app.utils.errors import AppError

FAKE_URI = "mongodb+srv://fake-user:fake-pass@fake-cluster.example.net/"
FAKE_GROQ_KEY = "gsk_fake_test_key_do_not_use"
FAKE_MODEL = "llama-3.3-70b-versatile"


class FakeAdmin:
    """Stands in for `client.admin`, the object the ping command runs on."""

    def __init__(self, should_fail: bool) -> None:
        self.should_fail = should_fail
        self.ping_count = 0

    def command(self, command_name: str) -> dict:
        self.ping_count += 1
        if self.should_fail:
            raise ConnectionFailure("simulated connection failure")
        return {"ok": 1.0}


class FakeDatabase(str):
    """Stands in for a pymongo Database.

    Subclasses str so existing assertions comparing it to
    "fake-database:<name>" still hold, while also supporting the
    `database[collection]` access that index creation needs.
    """

    def __getitem__(self, name):  # type: ignore[override]
        return FakeUsersCollection()


class FakeMongoClient:
    """Minimal stand-in for pymongo.MongoClient."""

    def __init__(self, should_fail: bool = False) -> None:
        self.admin = FakeAdmin(should_fail)
        self.closed = False

    def __getitem__(self, name: str) -> FakeDatabase:
        return FakeDatabase(f"fake-database:{name}")

    def close(self) -> None:
        self.closed = True


@pytest.fixture(autouse=True)
def reset_singletons() -> Iterator[None]:
    """Return the shared singletons to a clean state after each test.

    Without this, a test that fakes a connection would leak that state into
    the next test and make failures depend on ordering.
    """
    yield
    mongodb._client = None
    mongodb._connected = False
    groq_service._client = None


# --- Groq fakes -------------------------------------------------------------
#
# These mirror only the parts of the SDK response the service actually reads,
# so no test needs a real API key or a network connection.


class FakeMessage:
    def __init__(self, content: str | None) -> None:
        self.content = content


class FakeChoice:
    def __init__(self, content: str | None) -> None:
        self.message = FakeMessage(content)


class FakeUsage:
    def __init__(self, prompt: int, completion: int) -> None:
        self.prompt_tokens = prompt
        self.completion_tokens = completion
        self.total_tokens = prompt + completion


class FakeCompletion:
    """Stands in for the object returned by chat.completions.create()."""

    def __init__(
        self,
        content: str | None = "Hello from a fake model.",
        model: str = FAKE_MODEL,
        with_usage: bool = True,
        with_choices: bool = True,
    ) -> None:
        self.choices = [FakeChoice(content)] if with_choices else []
        self.model = model
        self.usage = FakeUsage(12, 8) if with_usage else None


class FakeCompletions:
    def __init__(self, result: object | Exception) -> None:
        self.result = result
        self.calls: list[dict] = []

    def create(self, **kwargs: object) -> object:
        self.calls.append(kwargs)
        if isinstance(self.result, Exception):
            raise self.result
        return self.result


class FakeGroqClient:
    """Minimal stand-in for groq.Groq."""

    def __init__(self, result: object | Exception) -> None:
        self.chat = type("FakeChat", (), {"completions": FakeCompletions(result)})()
        self.closed = False

    def close(self) -> None:
        self.closed = True


@pytest.fixture
def groq_configured(monkeypatch: pytest.MonkeyPatch) -> None:
    """Simulate GROQ_API_KEY being present, without a real key."""
    monkeypatch.setattr(settings, "GROQ_API_KEY", FAKE_GROQ_KEY)
    monkeypatch.setattr(settings, "GROQ_MODEL", FAKE_MODEL)
    groq_service._client = None


@pytest.fixture
def groq_unconfigured(monkeypatch: pytest.MonkeyPatch) -> None:
    """Simulate GROQ_API_KEY being absent."""
    monkeypatch.setattr(settings, "GROQ_API_KEY", None)
    groq_service._client = None


def install_fake_groq(result: object | Exception) -> FakeGroqClient:
    """Attach a fake client to the shared service and return it.

    Use together with the `groq_configured` fixture.
    """
    fake = FakeGroqClient(result)
    groq_service._client = fake
    return fake


# --- Agent-level fake -------------------------------------------------------
#
# Agent and route tests do not care how the SDK behaves — that is covered in
# test_groq_service.py. They replace GroqService.generate outright and inspect
# what the agent asked for.


class GenerateRecorder:
    """Records every call to groq_service.generate and returns a canned result."""

    def __init__(
        self,
        content: str = "Generated content.",
        error: Exception | None = None,
    ) -> None:
        self.content = content
        self.error = error
        self.calls: list[dict] = []

    def __call__(self, **kwargs: object) -> GroqResult:
        self.calls.append(kwargs)
        if self.error is not None:
            raise self.error
        return GroqResult(
            content=self.content,
            model=FAKE_MODEL,
            usage={"prompt_tokens": 30, "completion_tokens": 70, "total_tokens": 100},
        )

    @property
    def last(self) -> dict:
        """The keyword arguments of the most recent call."""
        return self.calls[-1]

    @property
    def system_prompt(self) -> str:
        return str(self.last.get("system_prompt", ""))

    @property
    def user_prompt(self) -> str:
        return str(self.last.get("user_prompt", ""))


@pytest.fixture
def recorded_generate(monkeypatch: pytest.MonkeyPatch) -> GenerateRecorder:
    """Replace the AI call with a recorder. No network, no API key needed."""
    recorder = GenerateRecorder()
    monkeypatch.setattr(groq_service, "generate", recorder)
    return recorder


# --- Document builders ------------------------------------------------------
#
# Real PDF and DOCX bytes, built in memory. No test depends on a checked-in
# binary fixture or on a document from the developer's machine.


def make_pdf(page_texts: list[str]) -> bytes:
    """Build a minimal but valid PDF with the given text on each page.

    Hand-built rather than generated by a library so the tests do not gain a
    dependency that production does not have. pypdf can create pages but
    cannot draw text, so a real text-bearing PDF has to be assembled here.
    """
    page_count = len(page_texts)
    page_ids = [4 + 2 * i for i in range(page_count)]
    content_ids = [5 + 2 * i for i in range(page_count)]

    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        (
            f"<< /Type /Pages /Kids "
            f"[{' '.join(f'{pid} 0 R' for pid in page_ids)}] /Count {page_count} >>"
        ).encode(),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    for index, text in enumerate(page_texts):
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 3 0 R >> >> "
                f"/Contents {content_ids[index]} 0 R >>"
            ).encode()
        )
        escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        stream = f"BT /F1 14 Tf 72 700 Td ({escaped}) Tj ET".encode()
        objects.append(
            b"<< /Length "
            + str(len(stream)).encode()
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_position = len(out)
    size = len(objects) + 1
    out += f"xref\n0 {size}\n".encode() + b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {size} /Root 1 0 R >>\n"
        f"startxref\n{xref_position}\n%%EOF\n"
    ).encode()

    return bytes(out)


def make_pdf_without_text(page_count: int = 1) -> bytes:
    """A structurally valid PDF whose pages carry no extractable text.

    Stands in for a scanned or image-only document.
    """
    from pypdf import PdfWriter

    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=612, height=792)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_docx(
    paragraphs: list[str] | None = None,
    tables: list[list[list[str]]] | None = None,
) -> bytes:
    """Build a real .docx with the given paragraphs and tables, in order."""
    from docx import Document

    document = Document()
    for text in paragraphs or []:
        document.add_paragraph(text)

    for table_rows in tables or []:
        if not table_rows:
            continue
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for cell_index, value in enumerate(row):
                table.cell(row_index, cell_index).text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_interleaved(blocks: list[tuple[str, object]]) -> bytes:
    """Build a .docx with paragraphs and tables in the given order.

    `make_docx` appends all paragraphs then all tables, which cannot exercise
    reading-order preservation. This builder emits blocks in sequence, so a
    table really does sit between two paragraphs in the body XML.

    blocks: [("p", "text"), ("table", [["a", "b"]]), ("p", "more")]
    """
    from docx import Document

    document = Document()
    for kind, value in blocks:
        if kind == "p":
            document.add_paragraph(str(value))
        elif kind == "table":
            rows = value  # type: ignore[assignment]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for row_index, row in enumerate(rows):
                for cell_index, cell_value in enumerate(row):
                    table.cell(row_index, cell_index).text = cell_value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- Users collection fake --------------------------------------------------
#
# An in-memory stand-in for the MongoDB users collection, including the unique
# email index. Written by hand rather than pulling in mongomock: it only needs
# find_one, insert_one and create_index, and the unique-index behaviour is
# precisely the part worth modelling explicitly.


class FakeUsersCollection:
    """Minimal users collection with a working unique constraint on email."""

    def __init__(self, fail: bool = False) -> None:
        self.documents: list[dict] = []
        self.fail = fail  # simulate MongoDB being unreachable
        self.indexes: list[tuple] = []

    def _check(self) -> None:
        if self.fail:
            from pymongo.errors import ServerSelectionTimeoutError

            raise ServerSelectionTimeoutError("simulated database outage")

    def create_index(self, key, **kwargs):
        self.indexes.append((key, kwargs))
        return kwargs.get("name", str(key))

    def find_one(self, query: dict) -> dict | None:
        self._check()
        for document in self.documents:
            if all(document.get(field) == value for field, value in query.items()):
                return dict(document)
        return None

    def insert_one(self, document: dict):
        self._check()

        # The real safeguard is the unique index, so the fake enforces it the
        # same way — by rejecting the write, not by checking beforehand.
        if any(d["email"] == document["email"] for d in self.documents):
            from pymongo.errors import DuplicateKeyError

            raise DuplicateKeyError("E11000 duplicate key error: email")

        from bson import ObjectId

        stored = dict(document)
        stored["_id"] = ObjectId()
        self.documents.append(stored)

        class Result:
            inserted_id = stored["_id"]

        return Result()


@pytest.fixture
def users(monkeypatch: pytest.MonkeyPatch) -> FakeUsersCollection:
    """Point the auth service at an in-memory users collection."""
    collection = FakeUsersCollection()
    monkeypatch.setattr(
        "app.services.auth_service.get_users_collection", lambda: collection
    )
    return collection


@pytest.fixture
def failing_users(monkeypatch: pytest.MonkeyPatch) -> FakeUsersCollection:
    """Simulate MongoDB being unreachable during an auth operation."""
    collection = FakeUsersCollection(fail=True)
    monkeypatch.setattr(
        "app.services.auth_service.get_users_collection", lambda: collection
    )
    return collection


@pytest.fixture
def jwt_secret(monkeypatch: pytest.MonkeyPatch) -> str:
    """A throwaway signing secret. Never a real one, and never from .env."""
    secret = "test-only-signing-secret-not-used-anywhere-real"
    monkeypatch.setattr(settings, "JWT_SECRET", secret)
    monkeypatch.setattr(settings, "JWT_ALGORITHM", "HS256")
    return secret


@pytest.fixture
def no_jwt_secret(monkeypatch: pytest.MonkeyPatch) -> None:
    """Simulate JWT_SECRET never having been configured."""
    monkeypatch.setattr(settings, "JWT_SECRET", None)


class FakeUpload:
    """Stand-in for FastAPI's UploadFile, for service-level tests."""

    def __init__(
        self, filename: str | None, content_type: str | None, data: bytes
    ) -> None:
        self.filename = filename
        self.content_type = content_type
        self._stream = io.BytesIO(data)

    async def read(self, size: int = -1) -> bytes:
        return self._stream.read(size)


@pytest.fixture
def failing_generate(monkeypatch: pytest.MonkeyPatch) -> GenerateRecorder:
    """Simulate the AI service raising a provider error."""
    recorder = GenerateRecorder(
        error=AppError(
            code="AI_PROVIDER_ERROR",
            message="AI service is temporarily unavailable",
            status_code=502,
        )
    )
    monkeypatch.setattr(groq_service, "generate", recorder)
    return recorder


@pytest.fixture
def unconfigured_db(monkeypatch: pytest.MonkeyPatch) -> None:
    """Simulate MONGODB_URI being absent."""
    monkeypatch.setattr(settings, "MONGODB_URI", None)
    mongodb._client = None
    mongodb._connected = False


@pytest.fixture
def connected_db(monkeypatch: pytest.MonkeyPatch) -> FakeMongoClient:
    """Simulate a configured database whose ping succeeds."""
    monkeypatch.setattr(settings, "MONGODB_URI", FAKE_URI)
    fake = FakeMongoClient(should_fail=False)
    mongodb._client = fake
    mongodb._connected = True
    return fake


@pytest.fixture
def failing_db(monkeypatch: pytest.MonkeyPatch) -> FakeMongoClient:
    """Simulate a configured database that cannot be reached."""
    monkeypatch.setattr(settings, "MONGODB_URI", FAKE_URI)
    fake = FakeMongoClient(should_fail=True)
    mongodb._client = fake
    mongodb._connected = True  # stale optimism; ping should correct it
    return fake
